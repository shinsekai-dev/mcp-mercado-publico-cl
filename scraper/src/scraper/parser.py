"""Parser de documentos para extraer texto."""

from pathlib import Path
from typing import Optional
import io

import fitz  # PyMuPDF
from docx import Document
from rich.console import Console

console = Console()


class DocumentParser:
    """Extrae texto de documentos PDF y Word."""

    @staticmethod
    def parse_pdf(filepath: Path) -> str:
        """Extrae texto de un archivo PDF."""
        try:
            text_parts = []
            doc = fitz.open(str(filepath))
            for page_num in range(len(doc)):
                page = doc[page_num]
                text = str(page.get_text())
                if text and text.strip():
                    text_parts.append(f"--- Página {page_num + 1} ---\n{text}")
            doc.close()

            return "\n\n".join(text_parts)
        except Exception as e:
            console.print(f"[red][ERROR][/red] Error parseando PDF {filepath}: {e}")
            return ""

    @staticmethod
    def parse_docx(filepath: Path) -> str:
        """Extrae texto de un archivo DOCX."""
        try:
            doc = Document(str(filepath))
            text_parts = []

            # Extraer de párrafos
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)

            # Extraer de tablas
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    if row_text:
                        text_parts.append(row_text)

            return "\n\n".join(text_parts)
        except Exception as e:
            console.print(f"[red][ERROR][/red] Error parseando DOCX {filepath}: {e}")
            return ""

    @staticmethod
    def parse_doc(filepath: Path) -> str:
        """
        Extrae texto de un archivo DOC (formato antiguo).
        Nota: python-docx no soporta .doc, necesitamos conversión o OCR.
        """
        console.print(
            f"[yellow][WARN][/yellow] Archivo .doc no soportado directamente: {filepath}"
        )
        console.print("[dim]   Considera instalar 'antiword' o convertir a .docx[/dim]")
        return f"[Archivo DOC no parseado: {filepath.name}]"

    @classmethod
    def parse_file(cls, filepath: Path) -> str:
        """
        Detecta el tipo de archivo y extrae texto.

        Args:
            filepath: Path al archivo

        Returns:
            Texto extraído o mensaje de error
        """
        if not filepath.exists():
            return f"[Archivo no encontrado: {filepath}]"

        suffix = filepath.suffix.lower()

        if suffix == ".pdf":
            return cls.parse_pdf(filepath)
        elif suffix == ".docx":
            return cls.parse_docx(filepath)
        elif suffix == ".doc":
            return cls.parse_doc(filepath)
        elif suffix in [".txt", ".md"]:
            return filepath.read_text(encoding="utf-8")
        else:
            console.print(f"[yellow][WARN][/yellow] Formato no soportado: {suffix}")
            return f"[Archivo no parseado: {filepath.name}]"

    @classmethod
    def parse_directory(
        cls, directory: Path, output_file: Optional[Path] = None
    ) -> str:
        """
        Parsea todos los documentos en un directorio y genera un resumen.

        Args:
            directory: Directorio con documentos
            output_file: Archivo opcional para guardar el resumen

        Returns:
            Texto completo de todos los documentos
        """
        if not directory.exists():
            return ""

        console.print(f"[blue][DOC] Parseando documentos en {directory}...[/blue]")

        all_texts = []
        document_files = sorted(directory.iterdir())

        for filepath in document_files:
            if filepath.is_file():
                console.print(f"[dim]  > {filepath.name}[/dim]")
                text = cls.parse_file(filepath)
                if text:
                    all_texts.append(f"\n{'=' * 60}\n")
                    all_texts.append(f"DOCUMENTO: {filepath.name}\n")
                    all_texts.append(f"{'=' * 60}\n\n")
                    all_texts.append(text)

        full_text = "\n\n".join(all_texts)

        if output_file:
            output_file.write_text(full_text, encoding="utf-8")
            console.print(f"[green][OK][/green] Resumen guardado en {output_file}")

        return full_text
