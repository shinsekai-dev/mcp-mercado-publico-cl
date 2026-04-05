"""Generador de documentos DOCX para postulación a licitaciones."""

from datetime import date
from pathlib import Path
from typing import Optional

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from domain.documentos.entities import TipoDocumento, CATALOGO_DOCUMENTOS
from infrastructure import brand


# ── Helpers de estilo ─────────────────────────────────────────────────────────

def _hex_to_rgb(hex_color: str) -> RGBColor:
    h = hex_color.lstrip("#")
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


def _set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _heading(doc: Document, text: str, level: int = 1):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text.upper())
    run.font.name = brand.FONT_PRIMARY
    run.font.size = Pt(brand.FONT_SIZE_TITLE if level == 1 else brand.FONT_SIZE_SUBTITLE)
    run.font.bold = True
    run.font.color.rgb = _hex_to_rgb(brand.COLOR_PRIMARY)
    return p


def _body(doc: Document, text: str, bold: bool = False, italic: bool = False, size: int = None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = brand.FONT_PRIMARY
    run.font.size = Pt(size or brand.FONT_SIZE_BODY)
    run.font.bold = bold
    run.font.italic = italic
    return p


def _field(doc: Document, label: str, value: str):
    p = doc.add_paragraph()
    run_label = p.add_run(f"{label}: ")
    run_label.font.name = brand.FONT_PRIMARY
    run_label.font.size = Pt(brand.FONT_SIZE_BODY)
    run_label.font.bold = True
    run_value = p.add_run(value)
    run_value.font.name = brand.FONT_PRIMARY
    run_value.font.size = Pt(brand.FONT_SIZE_BODY)
    return p


def _divider(doc: Document):
    p = doc.add_paragraph("─" * 80)
    p.runs[0].font.color.rgb = _hex_to_rgb(brand.COLOR_BORDER)
    p.runs[0].font.size = Pt(8)


def _firma(doc: Document, datos: dict, ciudad: str = "Santiago"):
    doc.add_paragraph()
    fecha = date.today().strftime("%d de %B de %Y").replace(
        "January", "enero").replace("February", "febrero").replace(
        "March", "marzo").replace("April", "abril").replace(
        "May", "mayo").replace("June", "junio").replace(
        "July", "julio").replace("August", "agosto").replace(
        "September", "septiembre").replace("October", "octubre").replace(
        "November", "noviembre").replace("December", "diciembre")
    _body(doc, f"{ciudad}, {fecha}")
    doc.add_paragraph()
    _body(doc, "_" * 40)
    _body(doc, datos.get("representante_legal", "___________________"), bold=True)
    _body(doc, "Representante Legal")
    _body(doc, datos.get("empresa", "___________________"))
    _body(doc, f"RUT: {datos.get('rut', '___________________')}")


def _header_empresa(doc: Document, datos: dict):
    """Bloque de encabezado con datos de la empresa."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    _set_cell_bg(cell, brand.COLOR_PRIMARY)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(datos.get("empresa", brand.EMPRESA).upper())
    run.font.name = brand.FONT_PRIMARY
    run.font.size = Pt(brand.FONT_SIZE_TITLE)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run(
        f"RUT: {datos.get('rut', brand.RUT)}  |  {datos.get('giro', brand.GIRO)}"
    )
    run2.font.name = brand.FONT_PRIMARY
    run2.font.size = Pt(brand.FONT_SIZE_SMALL)
    run2.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    doc.add_paragraph()


# ── Generadores de documentos específicos ────────────────────────────────────

def _generar_carta_presentacion(datos: dict, codigo: str, nombre_lic: str, organismo: str) -> Document:
    doc = Document()
    _set_margins(doc)
    _header_empresa(doc, datos)
    _heading(doc, "CARTA DE PRESENTACIÓN")
    _divider(doc)
    doc.add_paragraph()

    fecha = date.today().strftime("%d de %B de %Y").replace(
        "January", "enero").replace("February", "febrero").replace(
        "March", "marzo").replace("April", "abril").replace("May", "mayo").replace(
        "June", "junio").replace("July", "julio").replace("August", "agosto").replace(
        "September", "septiembre").replace("October", "octubre").replace(
        "November", "noviembre").replace("December", "diciembre")

    _body(doc, f"Santiago, {fecha}")
    doc.add_paragraph()
    _body(doc, f"Señores")
    _body(doc, organismo, bold=True)
    _body(doc, "Presente")
    doc.add_paragraph()
    _heading(doc, f"Ref.: Oferta Licitación {codigo}", level=2)
    doc.add_paragraph()

    empresa = datos.get("empresa", brand.EMPRESA)
    rut = datos.get("rut", brand.RUT)
    rep = datos.get("representante_legal", brand.REPRESENTANTE_LEGAL)

    texto = (
        f"Por medio de la presente, {empresa}, RUT {rut}, representada legalmente "
        f"por {rep}, tiene el agrado de presentar su oferta para la licitación "
        f"pública ID {codigo}, denominada «{nombre_lic}», convocada por {organismo}."
    )
    _body(doc, texto)
    doc.add_paragraph()
    _body(doc,
        "Declaramos conocer íntegramente las Bases Administrativas, Técnicas y "
        "Económicas de la presente licitación, y nos comprometemos a cumplir con "
        "todos los requisitos, plazos y condiciones establecidos en ellas."
    )
    doc.add_paragraph()
    _body(doc,
        "Asimismo, declaramos que los antecedentes presentados son fidedignos y "
        "que la empresa se encuentra habilitada para contratar con el Estado."
    )
    doc.add_paragraph()
    _body(doc, "Sin otro particular, saluda atentamente,")
    _firma(doc, datos)
    return doc


def _generar_anexo_aceptacion_bases(datos: dict, codigo: str, nombre_lic: str, organismo: str) -> Document:
    doc = Document()
    _set_margins(doc)
    _header_empresa(doc, datos)
    _heading(doc, "ANEXO N°2")
    _heading(doc, "DECLARACIÓN DE ACEPTACIÓN DE BASES", level=2)
    _divider(doc)
    doc.add_paragraph()

    empresa = datos.get("empresa", brand.EMPRESA)
    rut = datos.get("rut", brand.RUT)
    rep = datos.get("representante_legal", brand.REPRESENTANTE_LEGAL)

    _body(doc,
        f"Yo, {rep}, cédula de identidad N° ___________, en representación de "
        f"{empresa}, RUT {rut}, en calidad de Representante Legal, declaro bajo "
        f"juramento que:"
    )
    doc.add_paragraph()

    items = [
        f"He leído y comprendo en su totalidad las Bases Administrativas, Técnicas y Económicas de la Licitación Pública ID {codigo} — «{nombre_lic}».",
        f"Acepto expresamente todas las condiciones, plazos, requisitos y restricciones establecidas en dichas Bases.",
        "Reconozco que el incumplimiento de las condiciones establecidas podrá significar la eliminación de mi oferta.",
        "Acepto que los plazos son fatales e improrrogables salvo resolución fundada de la entidad licitante.",
    ]
    for i, item in enumerate(items, 1):
        p = doc.add_paragraph(style="List Number")
        run = p.add_run(item)
        run.font.name = brand.FONT_PRIMARY
        run.font.size = Pt(brand.FONT_SIZE_BODY)

    _firma(doc, datos)
    return doc


def _generar_anexo_conflicto_intereses(datos: dict, codigo: str, nombre_lic: str, organismo: str) -> Document:
    doc = Document()
    _set_margins(doc)
    _header_empresa(doc, datos)
    _heading(doc, "ANEXO N°3")
    _heading(doc, "DECLARACIÓN DE AUSENCIA DE CONFLICTO DE INTERESES", level=2)
    _divider(doc)
    doc.add_paragraph()

    empresa = datos.get("empresa", brand.EMPRESA)
    rut = datos.get("rut", brand.RUT)
    rep = datos.get("representante_legal", brand.REPRESENTANTE_LEGAL)

    _body(doc,
        f"Yo, {rep}, cédula de identidad N° ___________, en representación de "
        f"{empresa}, RUT {rut}, declaro bajo juramento que ni la empresa ni sus "
        f"representantes, directores o socios se encuentran en ninguna de las "
        f"situaciones de conflicto de interés descritas en el artículo 4° de la "
        f"Ley N° 19.886 respecto de {organismo}."
    )
    doc.add_paragraph()

    situaciones = [
        "No tengo relación de parentesco con funcionarios directivos del organismo licitante.",
        "La empresa no es una sociedad de personas en que formen parte funcionarios directivos del organismo licitante.",
        "No se han celebrado contratos de honorarios con funcionarios directivos del organismo licitante en los últimos dos años.",
        "No existe ninguna otra situación que genere conflicto de interés entre la empresa y el organismo convocante.",
    ]
    for item in situaciones:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(item)
        run.font.name = brand.FONT_PRIMARY
        run.font.size = Pt(brand.FONT_SIZE_BODY)

    _firma(doc, datos)
    return doc


def _generar_anexo_probidad(datos: dict, codigo: str, nombre_lic: str, organismo: str) -> Document:
    doc = Document()
    _set_margins(doc)
    _header_empresa(doc, datos)
    _heading(doc, "ANEXO N°4")
    _heading(doc, "DECLARACIÓN JURADA DE PROBIDAD", level=2)
    _divider(doc)
    doc.add_paragraph()

    empresa = datos.get("empresa", brand.EMPRESA)
    rut = datos.get("rut", brand.RUT)
    rep = datos.get("representante_legal", brand.REPRESENTANTE_LEGAL)

    _body(doc,
        f"Yo, {rep}, cédula de identidad N° ___________, Representante Legal de "
        f"{empresa}, RUT {rut}, declaro bajo juramento que la empresa y sus "
        f"representantes no se encuentran afectos a ninguna de las inhabilidades "
        f"contempladas en el artículo 4° de la Ley N° 19.886 de Bases sobre "
        f"Contratos Administrativos de Suministro y Prestación de Servicios."
    )
    doc.add_paragraph()

    inhabilidades = [
        "No haber sido condenado por prácticas antisindicales o infracción a los derechos fundamentales del trabajador.",
        "No haber sido condenado por delitos concursales, cohecho, lavado de activos u otros delitos de corrupción.",
        "No encontrarse en estado de quiebra o insolvencia.",
        "No haber sido declarado inhábil en el registro de proveedores del Estado (ChileProveedores).",
        "No tener deudas tributarias o previsionales en mora.",
    ]
    for item in inhabilidades:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(item)
        run.font.name = brand.FONT_PRIMARY
        run.font.size = Pt(brand.FONT_SIZE_BODY)

    _firma(doc, datos)
    return doc


def _generar_anexo_transferencia(datos: dict, codigo: str, nombre_lic: str, organismo: str) -> Document:
    doc = Document()
    _set_margins(doc)
    _header_empresa(doc, datos)
    _heading(doc, "ANEXO N°5")
    _heading(doc, "DATOS PARA TRANSFERENCIA ELECTRÓNICA", level=2)
    _divider(doc)
    doc.add_paragraph()

    _body(doc,
        "En caso de resultar adjudicado, los pagos deberán realizarse mediante "
        "transferencia electrónica a los siguientes datos bancarios:"
    )
    doc.add_paragraph()

    campos = [
        ("Razón Social", datos.get("empresa", brand.EMPRESA)),
        ("RUT", datos.get("rut", brand.RUT)),
        ("Banco", datos.get("banco", "___________________")),
        ("Tipo de Cuenta", datos.get("tipo_cuenta", "Cuenta Corriente")),
        ("N° de Cuenta", datos.get("numero_cuenta", "___________________")),
        ("Email para notificación", datos.get("email_transferencia", datos.get("email", brand.EMAIL))),
    ]

    table = doc.add_table(rows=len(campos), cols=2)
    table.style = "Table Grid"
    for i, (label, value) in enumerate(campos):
        row = table.rows[i]
        cell_label = row.cells[0]
        cell_value = row.cells[1]
        _set_cell_bg(cell_label, brand.COLOR_BG_ALT.lstrip("#") if brand.COLOR_BG_ALT.startswith("#") else brand.COLOR_BG_ALT)
        run_l = cell_label.paragraphs[0].add_run(label)
        run_l.font.bold = True
        run_l.font.name = brand.FONT_PRIMARY
        run_l.font.size = Pt(brand.FONT_SIZE_BODY)
        run_v = cell_value.paragraphs[0].add_run(value)
        run_v.font.name = brand.FONT_PRIMARY
        run_v.font.size = Pt(brand.FONT_SIZE_BODY)

    doc.add_paragraph()
    _body(doc,
        "Declaro que los datos indicados son verídicos y me hago responsable de "
        "cualquier error en la información proporcionada.",
        italic=True, size=brand.FONT_SIZE_SMALL
    )
    _firma(doc, datos)
    return doc


def _generar_pacto_integridad(datos: dict, codigo: str, nombre_lic: str, organismo: str) -> Document:
    doc = Document()
    _set_margins(doc)
    _header_empresa(doc, datos)
    _heading(doc, "ANEXO N°7")
    _heading(doc, "PACTO DE INTEGRIDAD", level=2)
    _divider(doc)
    doc.add_paragraph()

    empresa = datos.get("empresa", brand.EMPRESA)
    rut = datos.get("rut", brand.RUT)
    rep = datos.get("representante_legal", brand.REPRESENTANTE_LEGAL)

    _body(doc,
        f"El oferente {empresa}, RUT {rut}, representado por {rep}, declara que, "
        f"por sí mismo o a través de sus empleados, asesores, representantes, "
        f"directores u otras personas, NO ha ofrecido, dado, ni acordado dar a "
        f"ningún funcionario del organismo {organismo} u otros organismos del "
        f"Estado, dinero u otra remuneración, regalo, favor, préstamo u otra "
        f"ventaja o beneficio para sí u otras personas a cambio de acciones, "
        f"decisiones u omisiones de ese funcionario en relación con la presente "
        f"licitación (ID {codigo})."
    )
    doc.add_paragraph()
    _body(doc,
        "Asimismo, el oferente declara haber leído y conocer el contenido del "
        "artículo 100 de la Ley N° 18.575 Orgánica Constitucional de Bases "
        "Generales de la Administración del Estado.",
        size=brand.FONT_SIZE_SMALL
    )
    _firma(doc, datos)
    return doc


def _set_margins(doc: Document, margin_cm: float = 2.5):
    for section in doc.sections:
        section.top_margin = Cm(margin_cm)
        section.bottom_margin = Cm(margin_cm)
        section.left_margin = Cm(margin_cm)
        section.right_margin = Cm(margin_cm)


# ── Registry ──────────────────────────────────────────────────────────────────

_GENERATORS = {
    TipoDocumento.CARTA_PRESENTACION: _generar_carta_presentacion,
    TipoDocumento.ANEXO_2_ACEPTACION_BASES: _generar_anexo_aceptacion_bases,
    TipoDocumento.ANEXO_3_CONFLICTO_INTERESES: _generar_anexo_conflicto_intereses,
    TipoDocumento.ANEXO_4_DECLARACION_PROBIDAD: _generar_anexo_probidad,
    TipoDocumento.ANEXO_5_DATOS_TRANSFERENCIA: _generar_anexo_transferencia,
    TipoDocumento.ANEXO_7_PACTO_INTEGRIDAD: _generar_pacto_integridad,
}


class DocumentoGenerator:

    @staticmethod
    def generar(
        tipo: TipoDocumento,
        datos: dict,
        codigo: str,
        nombre_licitacion: str,
        organismo: str,
        output_path: Path,
    ) -> Path:
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)

        fn = _GENERATORS.get(tipo)
        if fn is None:
            raise ValueError(f"Tipo de documento no soportado: {tipo}")

        doc = fn(datos, codigo, nombre_licitacion, organismo)
        doc.save(output_path)
        return output_path

    @staticmethod
    def generar_todos(
        datos: dict,
        codigo: str,
        nombre_licitacion: str,
        organismo: str,
        output_dir: Path,
    ) -> list[Path]:
        output_dir = Path(output_dir)
        output_dir.mkdir(parents=True, exist_ok=True)
        generados = []
        for tipo in TipoDocumento:
            filename = f"{tipo.value}.docx"
            path = DocumentoGenerator.generar(
                tipo=tipo,
                datos=datos,
                codigo=codigo,
                nombre_licitacion=nombre_licitacion,
                organismo=organismo,
                output_path=output_dir / filename,
            )
            generados.append(path)
        return generados
